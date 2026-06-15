"""
DOCXImporter - Importador para arquivos Word (.docx).
======================================================
Usa python-docx para extrair texto de documentos Word.
Extrai parágrafos e texto de tabelas.
"""

from __future__ import annotations

from pathlib import Path
from typing import List

from .base_importer import BaseImporter, ImportResult, ImporterError
from ..utils.logger import get_logger


class DOCXImporter(BaseImporter):
    """
    Importador para arquivos Microsoft Word (.docx).
    
    Características:
    - Extrai texto de parágrafos
    - Extrai texto de tabelas
    - Preserva estrutura básica do documento
    """
    
    SUPPORTED_EXTENSIONS: List[str] = ['.docx']
    
    def __init__(self) -> None:
        super().__init__()
        self._logger = get_logger(__name__)
        self._docx_module = None
    
    def _ensure_docx(self) -> None:
        """Importa python-docx sob demanda."""
        if self._docx_module is None:
            try:
                import docx
                self._docx_module = docx
            except ImportError:
                raise ImporterError(
                    what="Biblioteca python-docx não encontrada.",
                    why="O python-docx é necessário para ler arquivos Word.",
                    how="Instale com: pip install python-docx"
                )
    
    def can_handle(self, file_path: str) -> bool:
        """
        Verifica se o arquivo é um documento Word.
        
        Args:
            file_path: Caminho do arquivo.
            
        Returns:
            True se for um arquivo .docx.
        """
        ext = self._get_file_extension(file_path)
        return ext in self.SUPPORTED_EXTENSIONS
    
    def extract(self, file_path: str) -> ImportResult:
        """
        Extrai texto de um arquivo DOCX.
        
        Args:
            file_path: Caminho do arquivo DOCX.
            
        Returns:
            ImportResult com o texto e metadados.
            
        Raises:
            ImporterError: Se houver erro na extração.
        """
        self._ensure_docx()
        path = self._validate_file_exists(file_path)
        warnings: List[str] = []
        
        # Verificar extensão (não suporta .doc antigo)
        if path.suffix.lower() == '.doc':
            raise ImporterError(
                what="Formato .doc (Word 97-2003) não é suportado.",
                why="O formato antigo do Word (.doc) requer bibliotecas adicionais.",
                how="Abra o arquivo no Word e salve como .docx (formato moderno)."
            )
        
        try:
            doc = self._docx_module.Document(path)
        except Exception as e:
            error_msg = str(e).lower()
            
            if "bad zip file" in error_msg or "not a valid" in error_msg:
                raise ImporterError(
                    what="O arquivo não parece ser um documento Word válido.",
                    why="O conteúdo não corresponde ao formato .docx esperado.",
                    how="Verifique se o arquivo é realmente um documento Word "
                        "e não está corrompido."
                )
            
            if "permission" in error_msg:
                raise ImporterError(
                    what="Sem permissão para ler o arquivo.",
                    why="O sistema operacional bloqueou o acesso.",
                    how="Verifique as permissões do arquivo ou feche o Word se estiver aberto."
                )
            
            self._logger.error(f"Erro ao abrir DOCX: {e}")
            raise ImporterError(
                what="Erro ao abrir o documento Word.",
                why=str(e),
                how="Verifique se o arquivo é um .docx válido e não está corrompido."
            )
        
        # Extrair texto dos parágrafos
        paragraphs_text = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs_text.append(text)
        
        # Extrair texto das tabelas
        tables_text = []
        table_count = 0
        
        for table in doc.tables:
            table_count += 1
            table_rows = []
            
            for row in table.rows:
                row_cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_cells.append(cell_text)
                
                if row_cells:
                    table_rows.append(" | ".join(row_cells))
            
            if table_rows:
                tables_text.append("\n".join(table_rows))
        
        if table_count > 0:
            warnings.append(
                f"Documento contém {table_count} tabela(s). "
                "O texto das tabelas foi extraído e separado por '|'."
            )
        
        # Combinar parágrafos e tabelas
        all_text_parts = paragraphs_text + tables_text
        text = "\n\n".join(all_text_parts)
        
        if not text.strip():
            warnings.append("O documento está vazio ou não contém texto extraível.")
        
        # Metadados
        words = text.split()
        metadata = {
            'paragraph_count': len(paragraphs_text),
            'table_count': table_count,
            'word_count': len(words),
            'char_count': len(text),
            'file_size': self._get_file_size(file_path),
            'file_size_formatted': self._format_file_size(self._get_file_size(file_path)),
        }
        
        self._logger.info(
            f"DOCX importado: {metadata['word_count']} palavras, "
            f"{metadata['paragraph_count']} parágrafos, "
            f"{table_count} tabelas"
        )
        
        return ImportResult(
            text=text,
            source_file=str(path),
            encoding="utf-8",  # python-docx retorna UTF-8
            warnings=warnings,
            metadata=metadata,
        )
